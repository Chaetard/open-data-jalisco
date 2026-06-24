# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 open-data-jalisco contributors

"""Round-trip tests for the DOCX and legacy-XLS extractors.

Fixtures are generated at runtime (python-docx / xlwt) so there are no binary
blobs in the repo. These guard the two things that actually break: (1) text +
tables are extracted in order, and (2) corrupt/renamed files degrade to
``no_text`` instead of raising (which would re-mark the doc ``failed``).
"""
from pathlib import Path

import docx
import xlwt

from open_data_jalisco.adapters.extraction import build_default_registry
from open_data_jalisco.adapters.extraction.docx import DocxTextExtractor
from open_data_jalisco.adapters.extraction.xls import XlsTextExtractor


def test_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    d = docx.Document()
    d.add_paragraph("Acta de sesión del ayuntamiento de Tala")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Concepto"
    table.cell(0, 1).text = "Monto"
    table.cell(1, 0).text = "Nómina"
    table.cell(1, 1).text = "123456"
    p = tmp_path / "acta.docx"
    d.save(str(p))

    out = DocxTextExtractor().extract(p)

    assert not out.needs_ocr
    assert out.pages == []  # docx has no fixed pages → flat chunking
    assert "Acta de sesión del ayuntamiento de Tala" in out.full_text
    assert "Concepto | Monto" in out.full_text
    assert "Nómina | 123456" in out.full_text
    assert out.metadata["table_count"] == 1


def test_xls_extracts_sheets_and_cells(tmp_path: Path) -> None:
    wb = xlwt.Workbook()
    ws = wb.add_sheet("Ingresos")
    for col, header in enumerate(("Mes", "Total")):
        ws.write(0, col, header)
    ws.write(1, 0, "Enero")
    ws.write(1, 1, 1000)
    p = tmp_path / "ingresos.xls"
    wb.save(str(p))

    out = XlsTextExtractor().extract(p)

    assert not out.needs_ocr
    assert len(out.pages) == 1  # one page per non-empty sheet, like xlsx
    assert "# Ingresos" in out.full_text
    assert "Mes | Total" in out.full_text
    assert "Enero | 1000" in out.full_text  # number 1000.0 trimmed to "1000"


def test_corrupt_files_degrade_to_no_text(tmp_path: Path) -> None:
    bad_docx = tmp_path / "bad.docx"
    bad_docx.write_bytes(b"this is not a real docx")
    out_d = DocxTextExtractor().extract(bad_docx)
    assert out_d.full_text == "" and not out_d.needs_ocr

    bad_xls = tmp_path / "bad.xls"
    bad_xls.write_bytes(b"this is not a real xls")
    out_x = XlsTextExtractor().extract(bad_xls)
    assert out_x.full_text == "" and not out_x.needs_ocr


def test_registry_routes_docx_and_xls() -> None:
    reg = build_default_registry()
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert isinstance(reg.find(docx_mime, "docx"), DocxTextExtractor)
    assert isinstance(reg.find("application/vnd.ms-excel", "xls"), XlsTextExtractor)
    # zip/octet-stream mime mis-sniffs still route by extension
    assert isinstance(reg.find("application/zip", "docx"), DocxTextExtractor)
    assert isinstance(reg.find("application/octet-stream", "xls"), XlsTextExtractor)
