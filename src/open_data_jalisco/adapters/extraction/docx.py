# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 open-data-jalisco contributors

"""DOCX text extractor backed by python-docx.

Walks the document body in order, emitting paragraph text and table rows
(cells joined by ``" | "``, matching the spreadsheet extractors). DOCX has no
fixed pagination — Word reflows pages per renderer — so we return a single flat
``full_text`` with ``pages=[]`` and let the chunker use its paragraph-aware
flat path (page numbers stay null, which is honest). A file that can't be
opened (renamed/corrupt) yields an empty document (``needs_ocr=False``) so the
pipeline records ``no_text`` instead of crashing the batch.
"""
from __future__ import annotations

from pathlib import Path

from ...ports.text_extractor import ExtractedDocument
from ...shared.logging import get_logger

logger = get_logger(__name__)

_DOCX_MIME = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_DOCX_EXTENSIONS = {"docx"}


class DocxTextExtractor:
    def can_handle(self, mime_type: str, extension: str) -> bool:
        # Match by extension too: a .docx is a zip, so magic-byte mime sniffing
        # often reports application/zip — the extension is the reliable signal.
        return (
            mime_type.lower() in _DOCX_MIME
            or extension.lower().lstrip(".") in _DOCX_EXTENSIONS
        )

    def extract(self, path: Path) -> ExtractedDocument:
        # Imported lazily so python-docx is only needed when a DOCX is actually
        # processed, not at every pipeline import.
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            logger.warning("docx.extract.open_error path=%s err=%r", path, e)
            return ExtractedDocument(
                full_text="",
                pages=[],
                needs_ocr=False,
                metadata={"error": f"open_error: {e!r}"},
            )

        blocks: list[str] = []
        para_count = 0
        table_count = 0
        p_tag, tbl_tag = qn("w:p"), qn("w:tbl")
        for child in doc.element.body.iterchildren():
            try:
                if child.tag == p_tag:
                    text = Paragraph(child, doc).text.strip()
                    if text:
                        blocks.append(text)
                        para_count += 1
                elif child.tag == tbl_tag:
                    rows = _table_rows(Table(child, doc))
                    if rows:
                        blocks.append("\n".join(rows))
                        table_count += 1
            except Exception as e:  # one malformed block must not kill the doc
                logger.warning("docx.extract.block_error path=%s err=%r", path, e)

        # Blank-line separated so the chunker's flat path splits on paragraphs.
        full_text = "\n\n".join(blocks)
        return ExtractedDocument(
            full_text=full_text,
            pages=[],
            needs_ocr=False,
            metadata={"paragraph_count": para_count, "table_count": table_count},
        )


def _table_rows(table) -> list[str]:
    """Each non-empty table row → cells joined by ' | '.

    python-docx repeats a horizontally-merged cell once per grid column it
    spans, so we collapse consecutive duplicate cells to avoid noise.
    """
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            t = cell.text.strip()
            if not t or (cells and t == cells[-1]):
                continue
            cells.append(t)
        if cells:
            rows.append(" | ".join(cells))
    return rows
